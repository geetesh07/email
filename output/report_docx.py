"""
report_docx.py — Generate a structured Word (.docx) report.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import List, Dict
import os


def _add_styled_heading(doc: Document, text: str, level: int = 1):
    """Add a styled heading to the document."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    return heading


def _add_table_with_style(doc: Document, headers: List[str], rows: List[List[str]]):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    return table


def generate_docx_report(
    output_path: str,
    people: List[Dict],
    specs: list,
    tagged_sentences: list,
    unresolved: list,
    timeline: list,
    title: str = "Engineering Email Intelligence Report",
):
    """
    Generate a Word document report with all extracted data.
    
    Args:
        output_path: Path to save the .docx file
        people: List of dicts {name, email, company, role, emails_sent, specs_mentioned}
        specs: List of SpecRecord objects
        tagged_sentences: List of TaggedSentence objects
        unresolved: List of TaggedSentence objects (unresolved items)
        timeline: List of dicts {date, sentence, mentioned_by}
        title: Report title
    """
    doc = Document()

    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    doc.add_paragraph()  # spacer

    # ── Section 1: People ──
    _add_styled_heading(doc, "People Involved", level=1)
    if people:
        headers = ["Name", "Email", "Company", "Role"]
        rows = [
            [p.get("name", ""), p.get("email", ""), p.get("company", ""), p.get("role", "")]
            for p in people
        ]
        _add_table_with_style(doc, headers, rows)
    else:
        doc.add_paragraph("No people data extracted.", style="List Bullet")
    doc.add_paragraph()

    # ── Section 2: Extracted Specs ──
    _add_styled_heading(doc, "Extracted Specifications", level=1)
    if specs:
        headers = ["Category", "Value", "Unit", "Mentioned By", "Date"]
        rows = [
            [s.category, s.value, s.unit, s.mentioned_by, s.date_str]
            for s in specs
        ]
        _add_table_with_style(doc, headers, rows)
    else:
        doc.add_paragraph("No specifications extracted.", style="List Bullet")
    doc.add_paragraph()

    # ── Section 3: Engineering Sentences ──
    _add_styled_heading(doc, "Engineering-Relevant Sentences", level=1)
    if tagged_sentences:
        for ts in tagged_sentences:
            para = doc.add_paragraph()
            run = para.add_run(f"[{', '.join(ts.trigger_keywords)}] ")
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run = para.add_run(ts.sentence)
            run.font.size = Pt(10)
            meta = para.add_run(f"\n    — {ts.mentioned_by}  |  {ts.date_str}")
            meta.font.size = Pt(8)
            meta.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    else:
        doc.add_paragraph("No engineering sentences flagged.", style="List Bullet")
    doc.add_paragraph()

    # ── Section 4: Unresolved Items ──
    _add_styled_heading(doc, "Open / Unconfirmed Items", level=1)
    if unresolved:
        for item in unresolved:
            para = doc.add_paragraph()
            run = para.add_run(f'"{item.sentence}"')
            run.font.italic = True
            run.font.size = Pt(10)
            meta = para.add_run(f"\n    — {item.mentioned_by}  |  {item.date_str}")
            meta.font.size = Pt(8)
            meta.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
    else:
        doc.add_paragraph("No unresolved items found.", style="List Bullet")
    doc.add_paragraph()

    # ── Section 5: Timeline ──
    _add_styled_heading(doc, "Timeline", level=1)
    if timeline:
        headers = ["Date", "Event / Sentence", "Mentioned By"]
        rows = [
            [t.get("date", ""), t.get("sentence", ""), t.get("mentioned_by", "")]
            for t in timeline
        ]
        _add_table_with_style(doc, headers, rows)
    else:
        doc.add_paragraph("No timeline events extracted.", style="List Bullet")

    # Save
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"  ✓ Word report saved: {output_path}")
    return output_path
